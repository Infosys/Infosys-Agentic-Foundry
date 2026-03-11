import os
import asyncpg
import numpy as np
from typing import List, Dict, Any, Optional
import logging
from io import BytesIO
import PyPDF2
from docx import Document
import fitz
from PIL import Image
from dotenv import load_dotenv

load_dotenv()

try:
    from rapidocr_onnxruntime import RapidOCR
    OCR_ENGINE = RapidOCR()
    logger = logging.getLogger(__name__)
except ImportError:
    OCR_ENGINE = None
    logger = logging.getLogger(__name__)
except Exception as e:
    OCR_ENGINE = None
    logger = logging.getLogger(__name__)

from utils.postgres_vector_store_jsonb import PostgresVectorStoreJSONB
from utils.remote_model_client import get_remote_models


class EmbeddingProcessor:
    
    def __init__(self, pool: asyncpg.Pool):
        self.pool = pool
        self.vector_store = PostgresVectorStoreJSONB(pool)
        model_server = os.getenv('MODEL_SERVER_URL', 'http://localhost:5000')
        self.embedding_model, _ = get_remote_models(model_server)
    
    async def process_and_store(
        self,
        kb_id: str,
        texts: List[str],
        metadata_list: Optional[List[dict]] = None,
        created_by: str = "system",
        filename: str = ""
    ):
        try:
            embeddings = self.embedding_model.encode(texts, convert_to_numpy=True)
            if not isinstance(embeddings, np.ndarray):
                embeddings = np.array(embeddings)
            
            await self.vector_store.store_embeddings_by_id(
                kb_id=kb_id,
                chunks=texts,
                embeddings=embeddings,
                metadata_list=metadata_list,
                created_by=created_by,
                filename=filename
            )
        except Exception as e:
            logger.error(f"Error processing embeddings for KB ID '{kb_id}': {e}", exc_info=True)
    
    async def process_document(
        self,
        kb_id: str,
        file_content: Dict[str, Any],
        created_by: str = "system"
    ):
        try:
            filename = file_content['filename']
            content = file_content['content']
            
            pages_data = await self._extract_text_with_pages(content, filename)
            
            if not pages_data:
                return
            
            file_extension = os.path.splitext(filename)[1].lower()
            total_pages = len(pages_data)
            file_size_kb = len(content) / 1024
            
            all_chunks = []
            all_metadata = []
            
            for page_data in pages_data:
                page_text = page_data['text']
                page_num = page_data['page_number']
                
                if not page_text:
                    continue
                
                page_chunks = self._chunk_text(page_text)
                
                for chunk_idx, chunk in enumerate(page_chunks):
                    all_chunks.append(chunk)
                    all_metadata.append({
                        'filename': filename,
                        'file_type': file_extension,
                        'file_size_kb': round(file_size_kb, 2),
                        'total_pages': total_pages,
                        'page_number': page_num,
                        'page_chunk_index': chunk_idx,
                        'total_page_chunks': len(page_chunks),
                        'chunk_index': len(all_chunks) - 1,
                        'total_chunks': 0
                    })
            
            total_chunks = len(all_chunks)
            for metadata in all_metadata:
                metadata['total_chunks'] = total_chunks
            
            if all_chunks:
                await self.process_and_store(
                    kb_id=kb_id,
                    texts=all_chunks,
                    metadata_list=all_metadata,
                    created_by=created_by,
                    filename=filename
                )
        except Exception as e:
            logger.error(f"Error processing document for KB ID '{kb_id}': {e}", exc_info=True)
    
    async def _extract_text_with_pages(self, content: bytes, filename: str) -> List[Dict[str, Any]]:
        ext = os.path.splitext(filename)[1].lower()
        
        try:
            if ext == '.pdf':
                return self._extract_pages_from_pdf_with_ocr(content)
            elif ext == '.docx':
                return self._extract_pages_from_docx_with_ocr(content)
            elif ext == '.txt':
                text = self._extract_text_from_txt(content)
                return [{'text': text, 'page_number': 1}]
            elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
                text = self._extract_text_from_image(content)
                return [{'text': text, 'page_number': 1}]
            else:
                text = self._extract_text_from_txt(content)
                return [{'text': text, 'page_number': 1}]
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise
    
    def _extract_pages_from_pdf_with_ocr(self, content: bytes) -> List[Dict[str, Any]]:
        try:
            pages_data = []
            pdf_document = fitz.open(stream=content, filetype="pdf")
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text_blocks = page.get_text("blocks")
                image_list = page.get_images(full=True)
                content_items = []
                
                for block in text_blocks:
                    if len(block) >= 7 and block[6] == 0:
                        x0, y0, x1, y1, text, block_no, block_type = block[:7]
                        if text.strip():
                            content_items.append({
                                'type': 'text',
                                'position': (y0, x0),
                                'content': text.strip()
                            })
                
                if OCR_ENGINE and image_list:
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = pdf_document.extract_image(xref)
                            image_bytes = base_image["image"]
                            
                            img_rects = page.get_image_rects(xref)
                            if img_rects:
                                rect = img_rects[0]
                                y_pos = rect.y0
                                x_pos = rect.x0
                            else:
                                y_pos = float('inf')
                                x_pos = float('inf')
                            
                            image = Image.open(BytesIO(image_bytes))
                            result, _ = OCR_ENGINE(np.array(image))
                            
                            if result:
                                ocr_text = '\n'.join([item[1] for item in result])
                                if ocr_text.strip():
                                    content_items.append({
                                        'type': 'image',
                                        'position': (y_pos, x_pos),
                                        'content': f"[Image {img_index + 1}]\n{ocr_text.strip()}"
                                    })
                        except Exception as e:
                            continue
                
                content_items.sort(key=lambda x: (x['position'][0], x['position'][1]))
                
                if content_items:
                    combined_text = '\n\n'.join([item['content'] for item in content_items])
                    pages_data.append({
                        'text': combined_text,
                        'page_number': page_num + 1
                    })
            
            pdf_document.close()
            
            if not pages_data:
                return self._extract_pages_from_pdf_basic(content)
            
            return pages_data
        except Exception as e:
            return self._extract_pages_from_pdf_basic(content)
    
    def _extract_pages_from_pdf_basic(self, content: bytes) -> List[Dict[str, Any]]:
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            pages_data = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        pages_data.append({
                            'text': page_text.strip(),
                            'page_number': page_num + 1
                        })
                except Exception as e:
                    continue
            
            return pages_data
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise
    
    def _extract_pages_from_docx_with_ocr(self, content: bytes) -> List[Dict[str, Any]]:
        try:
            docx_file = BytesIO(content)
            doc = Document(docx_file)
            
            image_map = {}
            if OCR_ENGINE:
                for idx, rel in enumerate(doc.part.rels.values()):
                    if "image" in rel.reltype:
                        try:
                            image_data = rel.target_part.blob
                            image = Image.open(BytesIO(image_data))
                            result, _ = OCR_ENGINE(np.array(image))
                            
                            if result:
                                ocr_text = '\n'.join([item[1] for item in result])
                                if ocr_text.strip():
                                    image_map[rel.target_ref] = f"[Image {idx + 1}]\n{ocr_text.strip()}"
                        except Exception as e:
                            continue
            
            content_parts = []
            
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if 'graphicData' in run._element.xml:
                        for rel in run.element.xpath('.//a:blip/@r:embed'):
                            if rel in image_map:
                                content_parts.append(image_map[rel])
                                del image_map[rel]
                
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                if table_text:
                    content_parts.append('\n'.join(table_text))
            
            for image_text in image_map.values():
                content_parts.append(image_text)
            
            combined_text = '\n\n'.join(content_parts)
            return [{'text': combined_text, 'page_number': 1}] if combined_text else []
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise
    
    def _extract_text_from_image(self, content: bytes) -> str:
        try:
            image = Image.open(BytesIO(content))
            if image.mode in ('RGBA', 'P'):
                image = image.convert('RGB')
            
            if OCR_ENGINE:
                result, _ = OCR_ENGINE(np.array(image))
                if result:
                    return '\n'.join([item[1] for item in result])
            return ""
        except Exception as e:
            logger.error(f"Error performing OCR on image: {e}")
            raise
    
    def _extract_text_from_txt(self, content: bytes) -> str:
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error reading TXT: {e}")
            raise
    
    def _chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> List[str]:
        if not text:
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            
            if end < text_length:
                search_start = max(start, end - 200)
                sentence_end = max(
                    text.rfind('.', search_start, end),
                    text.rfind('!', search_start, end),
                    text.rfind('?', search_start, end)
                )
                
                if sentence_end > start:
                    end = sentence_end + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks
